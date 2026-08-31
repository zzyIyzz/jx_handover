"""V0.3.0 migration, chapters, preview-import and Word regression tests."""
from __future__ import annotations

import json
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
import sys
import tempfile
import threading
import unittest
from unittest import mock

from fastapi import HTTPException
from openpyxl import Workbook
from openpyxl.styles import Font
from sqlalchemy import create_engine, inspect, text
from sqlalchemy.orm import sessionmaker

PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "backend"))

from app import config
from app.db import Base
from app.migrations import migrate_database
from app.models import (
    DocumentSnapshot,
    ExternalAssessment,
    HandoverBatch,
    HandoverItem,
    HandoverStationMeta,
    ImportJob,
    MonthlyPlanItem,
    SectionImportPreview,
    SourceRecord,
    Station,
    WorkItem,
)
from app.services import handover_service, rules
from app.services.document import mapper, renderer, validator
from app.services.importer import sections
import launcher


class V030LauncherTest(unittest.TestCase):
    def test_fixed_local_port_and_unrelated_occupant_are_distinguished(self):
        self.assertEqual(launcher.HOST, "127.0.0.1")
        self.assertEqual(launcher.PORT, 8765)
        self.assertEqual(launcher.URL, "http://127.0.0.1:8765")

        class UnrelatedHandler(BaseHTTPRequestHandler):
            def do_GET(self):  # noqa: N802
                body = json.dumps({
                    "status": "ok",
                    "service": "another-program",
                    "port": 8765,
                }).encode("utf-8")
                self.send_response(200)
                self.send_header("Content-Type", "application/json")
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)

            def log_message(self, *_args):
                return

        try:
            server = ThreadingHTTPServer((launcher.HOST, launcher.PORT), UnrelatedHandler)
        except OSError as exc:
            self.skipTest(f"验收端口当前不可用：{exc}")
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        try:
            self.assertTrue(launcher.port_is_in_use())
            self.assertIsNone(launcher.health_payload())
        finally:
            server.shutdown()
            server.server_close()
            thread.join(timeout=3)

    def test_legacy_migration_backs_up_database_and_copies_user_files(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            legacy = root / "旧版项目" / "runtime"
            target = root / "LocalAppData" / "JXHandover"
            (legacy / "data").mkdir(parents=True)
            (legacy / "imports").mkdir()
            (legacy / "generated").mkdir()
            (legacy / "snapshots").mkdir()
            (legacy / "data" / "handover.db").write_bytes(b"legacy-database")
            (legacy / "imports" / "原件.xlsx").write_bytes(b"xlsx")
            (legacy / "generated" / "历史.docx").write_bytes(b"docx")
            (legacy / "snapshots" / "历史.json").write_text("{}", encoding="utf-8")

            replacements = {
                "USER_DATA_ROOT": target,
                "DATA_DIR": target / "data",
                "IMPORT_DIR": target / "imports",
                "GENERATED_DIR": target / "generated",
                "SNAPSHOT_DIR": target / "snapshots",
                "LOG_DIR": target / "logs",
                "DATABASE_PATH": target / "data" / "handover.db",
            }
            with mock.patch.multiple(config, **replacements):
                result = launcher.migrate_legacy_data(legacy)

            self.assertEqual(
                (target / "data" / "handover.db").read_bytes(),
                b"legacy-database",
            )
            self.assertTrue((target / "imports" / "原件.xlsx").exists())
            self.assertTrue((target / "generated" / "历史.docx").exists())
            self.assertTrue((target / "snapshots" / "历史.json").exists())
            backup = Path(result["legacy_database_backup"])
            self.assertTrue(backup.exists())
            self.assertEqual(backup.read_bytes(), b"legacy-database")
            self.assertEqual(len(list((target / "snapshots").glob("legacy_migration_*.json"))), 1)


class V030DatabaseTest(unittest.TestCase):
    def test_migration_backs_up_once_is_idempotent_and_preserves_history(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            database = root / "old.db"
            engine = create_engine(f"sqlite:///{database.as_posix()}")
            Base.metadata.create_all(engine)
            session = sessionmaker(bind=engine)()
            station = Station(code="TEST", name="测试场站", aliases_json='["测试"]')
            session.add(station)
            session.flush()
            batch = HandoverBatch(
                start_date="2026-08-01",
                end_date="2026-08-10",
                handover_date="2026-08-10",
            )
            session.add(batch)
            session.flush()
            meta = HandoverStationMeta(batch_id=batch.id, station_id=station.id)
            session.add(meta)
            session.flush()
            work = WorkItem(
                station_id=station.id,
                canonical_title="旧版已完成事项",
                status="completed",
                priority="normal",
                first_seen_date="2026-08-01",
                last_seen_date="2026-08-02",
            )
            session.add(work)
            session.flush()
            item = HandoverItem(
                batch_id=batch.id,
                station_meta_id=meta.id,
                work_item_id=work.id,
                title_snapshot=work.canonical_title,
                status="completed",
                priority="normal",
            )
            session.add(item)
            plan = MonthlyPlanItem(plan_month="2026-08", title="旧定期工作")
            session.add(plan)
            snapshot = DocumentSnapshot(
                batch_id=batch.id,
                station_meta_id=meta.id,
                version=1,
                status="published",
                data_json='{"legacy": true}',
                docx_path="history.docx",
                sha256="abc",
            )
            session.add(snapshot)
            session.commit()
            item_id = item.id
            snapshot_id = snapshot.id
            session.close()

            with engine.begin() as connection:
                connection.execute(text("DROP TABLE section_import_previews"))
                connection.execute(text("DROP TABLE external_assessments"))
                connection.execute(text("ALTER TABLE handover_items DROP COLUMN section"))
                connection.execute(text("ALTER TABLE handover_items DROP COLUMN completed_by"))
                connection.execute(text("ALTER TABLE handover_items DROP COLUMN sort_order"))
                connection.execute(text("ALTER TABLE import_jobs DROP COLUMN stored_path"))
                connection.execute(text("ALTER TABLE import_jobs DROP COLUMN parser_key"))
                connection.execute(text("ALTER TABLE monthly_plan_items DROP COLUMN library_id"))

            backup_dir = root / "backups"
            first = migrate_database(engine, backup_dir=backup_dir)
            self.assertIsNotNone(first["backup_path"])
            self.assertTrue(Path(first["backup_path"]).exists())
            columns = {column["name"] for column in inspect(engine).get_columns("handover_items")}
            self.assertTrue({"section", "completed_by", "sort_order"}.issubset(columns))
            with engine.connect() as connection:
                migrated = connection.execute(text(
                    "SELECT section, completed_by FROM handover_items WHERE id=:id"
                ), {"id": item_id}).one()
                history = connection.execute(text(
                    "SELECT data_json, docx_path FROM document_snapshots WHERE id=:id"
                ), {"id": snapshot_id}).one()
            self.assertEqual(migrated.section, "important")
            self.assertEqual(migrated.completed_by, "")
            self.assertEqual(history.data_json, '{"legacy": true}')
            self.assertEqual(history.docx_path, "history.docx")

            backup_count = len(list(backup_dir.glob("*.db")))
            second = migrate_database(engine, backup_dir=backup_dir)
            self.assertEqual(second["changed"], [])
            self.assertIsNone(second["backup_path"])
            self.assertEqual(len(list(backup_dir.glob("*.db"))), backup_count)
            engine.dispose()


class V030WorkflowTest(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(self.engine)
        self.db = sessionmaker(bind=self.engine)()
        self.station = Station(code="TEST", name="测试场站", aliases_json='["测试"]')
        self.db.add(self.station)
        self.db.flush()
        self.batch = HandoverBatch(
            start_date="2026-08-14",
            end_date="2026-08-23",
            handover_date="2026-08-23",
            status="review",
        )
        self.db.add(self.batch)
        self.db.flush()
        self.meta = HandoverStationMeta(
            batch_id=self.batch.id,
            station_id=self.station.id,
            duty_leader="甲",
            temp_leader="无",
            operators_json='["乙"]',
        )
        self.db.add(self.meta)
        self.db.commit()

    def tearDown(self) -> None:
        self.db.close()

    def _add_item(self, **values):
        payload = {
            "station_meta_id": self.meta.id,
            "title_snapshot": "测试事项",
            "status": "in_progress",
            "priority": "normal",
            "section": "handover",
        }
        payload.update(values)
        result = handover_service.add_handover_item(self.db, self.batch.id, payload)
        return self.db.get(HandoverItem, result["id"])

    def test_item_section_priority_sort_and_conflict_are_independent(self):
        completed = self._add_item(
            title_snapshot="已完成工作",
            status="completed",
            priority="urgent",
            section=None,
            completed_by="乙",
        )
        self.assertEqual(completed.section, "important")
        self.assertEqual(rules.professional_color(completed.priority, completed.status), "red")
        moved = handover_service.patch_item(
            self.db,
            completed.id,
            completed.revision,
            {"section": "handover"},
        )
        self.db.refresh(completed)
        self.assertEqual(completed.section, "handover")
        self.assertEqual(completed.priority, "urgent")
        with self.assertRaises(HTTPException) as conflict:
            handover_service.delete_handover_item(self.db, completed.id, moved["revision"] - 1)
        self.assertEqual(conflict.exception.status_code, 409)

        second = self._add_item(title_snapshot="第二项", sort_order=50)
        handover_service.reorder_handover_items(
            self.db,
            self.batch.id,
            self.meta.id,
            "handover",
            [second.id, completed.id],
        )
        self.db.refresh(second)
        self.db.refresh(completed)
        self.assertLess(second.sort_order, completed.sort_order)

    def test_external_assessment_crud_sort_and_conflict(self):
        first = handover_service.add_external_assessment(self.db, self.batch.id, {
            "station_meta_id": self.meta.id,
            "contractor": "甲单位",
            "work_content": "检修",
            "assessment": "合格",
            "remark": "",
        })
        second = handover_service.add_external_assessment(self.db, self.batch.id, {
            "station_meta_id": self.meta.id,
            "contractor": "乙单位",
            "work_content": "清扫",
            "assessment": "整改",
            "remark": "复查",
        })
        handover_service.reorder_external_assessments(
            self.db, self.batch.id, self.meta.id, [second["id"], first["id"]]
        )
        first_row = self.db.get(ExternalAssessment, first["id"])
        second_row = self.db.get(ExternalAssessment, second["id"])
        self.assertLess(second_row.sort_order, first_row.sort_order)
        with self.assertRaises(HTTPException) as conflict:
            handover_service.patch_external_assessment(
                self.db, first_row.id, first_row.revision - 1, {"assessment": "优秀"}
            )
        self.assertEqual(conflict.exception.status_code, 409)

    @staticmethod
    def _standard_workbook(path: Path) -> None:
        book = Workbook()
        book.remove(book.active)
        third = book.create_sheet("第三章-重点工作")
        third.append(["工作内容", "开始时间", "结束时间", "完成情况", "优先级", "完成人", "备注"])
        third.append(["重点试验", "2026-08-14", "2026-08-15", "已完成", "重点", "乙", "数据正常"])
        fourth = book.create_sheet("第四章-需交接")
        fourth.append(["工作内容", "开始时间", "结束时间", "完成情况", "优先级", "交接前责任人", "交接后责任人", "备注"])
        fourth.append(["整改事项", "2026-08-16", "2026-09-30", "进行中", "普通", "甲", "乙", "材料到场"])
        fifth = book.create_sheet("第五章-外委考核")
        fifth.append(["外委单位", "工作内容", "考核情况", "备注"])
        fifth.append(["测试外委", "设备检修", "合格", ""])
        book.save(path)

    def test_standard_preview_is_read_only_then_commit_is_idempotent(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            workbook = root / "标准导入.xlsx"
            self._standard_workbook(workbook)
            old_import_dir = config.IMPORT_DIR
            config.IMPORT_DIR = root / "archive"
            try:
                preview = sections.create_preview(self.db, self.batch.id, self.meta.id, workbook)
                self.assertEqual(preview["parser_key"], "standard_template")
                self.assertEqual(self.db.query(HandoverItem).count(), 0)
                self.assertEqual(self.db.query(ExternalAssessment).count(), 0)
                first = sections.commit_preview(self.db, self.batch.id, preview["id"])
                second = sections.commit_preview(self.db, self.batch.id, preview["id"])
                self.assertEqual(first, second)
                self.assertEqual(first["created_items"], 2)
                self.assertEqual(first["created_external_assessments"], 1)
                self.assertEqual(self.db.query(HandoverItem).count(), 2)
                self.assertEqual(self.db.query(SourceRecord).count(), 2)

                duplicate_preview = sections.create_preview(
                    self.db, self.batch.id, self.meta.id, workbook
                )
                self.assertEqual(duplicate_preview["summary"]["duplicate"], 3)
            finally:
                config.IMPORT_DIR = old_import_dir

    def test_work_log_adapter_forward_fills_and_uses_latest_status(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            workbook_path = root / "脱敏工作日志.xlsx"
            book = Workbook()
            sheet = book.active
            sheet.title = "Sheet1"
            sheet.append(["脱敏工作日志"])
            sheet.append([
                "date", "station", "工作内容", "完成期限", "工作要求",
                "人员安排", "是否完成", "是否列入本日计划", "工作进度",
            ])
            sheet.append([
                "2026-08-14", "测试场站", "F01试验", "2026-08-20", "按方案执行",
                "乙", "未完成", "是", "准备中",
            ])
            sheet["C3"].font = Font(color="FFFF0000")
            sheet.append([None, None, "排水沟整改", "2026-09-30", "完成整改", "甲", "未完成", "是", "材料到场"])
            sheet.merge_cells("A3:A4")
            sheet.merge_cells("B3:B4")
            sheet.append(["2026-08-15", "测试场站", "F01试验", "2026-08-20", "按方案执行", "乙", "已完成", "是", "试验正常"])
            sheet["C5"].font = Font(color="FFFF0000")
            book.create_sheet("Sheet2")
            book.create_sheet("Sheet3")
            book.save(workbook_path)

            old_import_dir = config.IMPORT_DIR
            config.IMPORT_DIR = root / "archive"
            try:
                preview = sections.create_preview(
                    self.db, self.batch.id, self.meta.id, workbook_path
                )
            finally:
                config.IMPORT_DIR = old_import_dir
            self.assertEqual(preview["parser_key"], "work_log")
            self.assertEqual(preview["summary"]["total"], 2)
            rows = {row["title_snapshot"]: row for row in preview["rows"]}
            self.assertEqual(rows["F01试验"]["status"], "completed")
            self.assertEqual(rows["F01试验"]["section"], "important")
            self.assertEqual(rows["F01试验"]["priority"], "urgent")
            self.assertEqual(rows["F01试验"]["start_date"], "2026-08-14")
            self.assertEqual(rows["F01试验"]["end_date"], "2026-08-15")
            self.assertEqual(rows["排水沟整改"]["section"], "handover")

    def test_mapper_renderer_preserve_sections_counts_colors_and_deduplicate_remark(self):
        important = self._add_item(
            title_snapshot="重点试验",
            status="completed",
            priority="important",
            section="important",
            completed_by="乙",
            summary="重点试验",
            latest_progress="重点试验",
            start_date="2026-08-14",
            end_date="2026-08-15",
        )
        handover = self._add_item(
            title_snapshot="排水沟整改",
            status="in_progress",
            priority="urgent",
            section="handover",
            previous_owner="甲",
            next_owner="乙",
            latest_progress="材料到场",
            start_date="2026-08-16",
            end_date="2026-09-30",
        )
        handover_service.add_external_assessment(self.db, self.batch.id, {
            "station_meta_id": self.meta.id,
            "contractor": "测试外委",
            "work_content": "设备检修",
            "assessment": "合格",
            "remark": "",
        })
        data = mapper.build_context(self.db, self.batch, self.meta)
        self.assertEqual(data["ctx"]["important_items"][0]["remark"], "")
        self.assertEqual(data["colors"]["important"], ["yellow"])
        self.assertEqual(data["colors"]["handover"], ["red"])

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "rendered.docx"
            template = Path(__file__).resolve().parents[1] / "app" / "templates" / "word" / "handover_v1.docx"
            renderer.render_word(template, data["ctx"], data["colors"], output)
            report = validator.validate_docx(output, data["expected"])
            self.assertTrue(report["valid"])
            from docx import Document
            from docx.oxml.ns import qn

            doc = Document(output)
            self.assertEqual(len(doc.tables), 7)
            self.assertEqual(doc.tables[1].rows[1].cells[1].text, important.title_snapshot)
            self.assertEqual(doc.tables[2].rows[1].cells[1].text, handover.title_snapshot)
            yellow = doc.tables[1].rows[1].cells[0]._tc.get_or_add_tcPr().find(qn("w:shd"))
            red = doc.tables[2].rows[1].cells[0]._tc.get_or_add_tcPr().find(qn("w:shd"))
            self.assertEqual(yellow.get(qn("w:fill")), "FFFE83")
            self.assertEqual(red.get(qn("w:fill")), "FFA5A5")
            for table in doc.tables[1:]:
                for row in table.rows:
                    self.assertIsNotNone(
                        row._tr.get_or_add_trPr().find(qn("w:cantSplit"))
                    )


if __name__ == "__main__":
    unittest.main()
