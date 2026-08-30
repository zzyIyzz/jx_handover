"""Render a handover DOCX by cloning the sanitized reference-table rows."""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph


TABLE_IMPORTANT = 1
TABLE_HANDOVER = 2
TABLE_EXTERNAL = 3
TABLE_MONTHLY = 4
TABLE_QUARTERLY = 5
TABLE_YEARLY = 6

ROW_COLORS = {
    "red": "FFA5A5",
    "yellow": "FFFE83",
    "green": "C6EFCE",
    "white": "FFFFFF",
}


def _replace_paragraph_text(paragraph: Paragraph, value: str) -> None:
    """Set text while retaining the first run's full reference formatting."""
    if paragraph.runs:
        keep = paragraph.runs[0]._r
        for run in paragraph.runs[1:]:
            paragraph._p.remove(run._r)
    else:
        keep = paragraph.add_run()._r
    for child in list(keep):
        if child.tag != qn("w:rPr"):
            keep.remove(child)
    lines = str(value or "").split("\n")
    for index, line in enumerate(lines):
        text = OxmlElement("w:t")
        if line.startswith(" ") or line.endswith(" "):
            text.set(qn("xml:space"), "preserve")
        text.text = line
        keep.append(text)
        if index < len(lines) - 1:
            keep.append(OxmlElement("w:br"))


def _set_cell_text(cell: _Cell, value: object) -> None:
    paragraphs = cell.paragraphs
    paragraph = paragraphs[0] if paragraphs else cell.add_paragraph()
    for extra in paragraphs[1:]:
        cell._tc.remove(extra._p)
    _replace_paragraph_text(paragraph, "" if value is None else str(value))


def _shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:fill"), fill)


def _style_row(row: _Row, color: str) -> None:
    fill = ROW_COLORS.get(color, ROW_COLORS["white"])
    for cell in row.cells:
        _shade_cell(cell, fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)


def _repeat_header(table: Table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def _populate_table(
    table: Table,
    rows: list[dict],
    fields: tuple[str, ...],
    colors: list[str] | None = None,
    *,
    minimum_rows: int = 1,
    empty_title: str = "",
) -> None:
    if len(table.rows) < 2:
        raise RuntimeError("Word template table has no model row")
    model_xml = deepcopy(table.rows[1]._tr)
    for existing in list(table.rows[1:]):
        table._tbl.remove(existing._tr)

    output_count = max(minimum_rows, len(rows))
    for index in range(output_count):
        table._tbl.append(deepcopy(model_xml))
        target = table.rows[-1]
        row_data = rows[index] if index < len(rows) else {}
        number = index + 1 if rows or minimum_rows > 1 or empty_title else ""
        _set_cell_text(target.cells[0], number)
        for column, field in enumerate(fields, start=1):
            value = row_data.get(field, "")
            if not rows and index == 0 and field == "title" and empty_title:
                value = empty_title
            _set_cell_text(target.cells[column], value)
        color = colors[index] if colors and index < len(colors) else "white"
        _style_row(target, color)
    _repeat_header(table)


def _set_basic_info(doc, context: dict) -> None:
    values = (
        context["start_date_cn"],
        context["end_date_cn"],
        context["handover_date_cn"],
        context["duty_leader"],
        context["temp_leader"],
        context["operators"],
    )
    table = doc.tables[0]
    if len(table.rows) != len(values):
        raise RuntimeError("Basic-information table no longer matches the reference")
    for row, value in zip(table.rows, values):
        _set_cell_text(row.cells[1], value)


def _set_title_and_devices(doc, context: dict) -> None:
    title = f"{context['station_name']}交接班记录\n（{context['period_cn']}班次）"
    _replace_paragraph_text(doc.paragraphs[0], title)
    heading_index = next(
        index for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.startswith("二、")
    )
    devices = context.get("device_changes") or []
    device_text = (
        "\n".join(f"{index}、{value}" for index, value in enumerate(devices, 1))
        if devices else "本班无设备变更。"
    )
    _replace_paragraph_text(doc.paragraphs[heading_index + 1], device_text)


def _protect_section_titles(doc) -> None:
    """Keep chapter/subchapter titles with their instruction or table."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if (
            text.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "6."))
            or text.startswith(("（", "("))
        ):
            paragraph.paragraph_format.keep_with_next = True


def render_word(
    template_path: Path,
    context: dict,
    colors: dict,
    output_path: Path,
) -> Path:
    doc = Document(str(template_path))
    if len(doc.tables) != 7:
        raise RuntimeError(f"V0.3 Word template must contain seven tables, got {len(doc.tables)}")

    _set_title_and_devices(doc, context)
    _protect_section_titles(doc)
    _set_basic_info(doc, context)
    _populate_table(
        doc.tables[TABLE_IMPORTANT],
        context.get("important_items", []),
        ("title", "start", "end", "completed_by", "remark"),
        colors.get("important", []),
        empty_title="本班无紧急/重点工作",
    )
    _populate_table(
        doc.tables[TABLE_HANDOVER],
        context.get("handover_items", []),
        (
            "title",
            "start",
            "end",
            "previous_owner",
            "next_owner",
            "status_text",
            "remark",
        ),
        colors.get("handover", []),
    )
    _populate_table(
        doc.tables[TABLE_EXTERNAL],
        context.get("external_rows", []),
        ("contractor", "work_content", "assessment", "remark"),
        minimum_rows=3 if not context.get("external_rows") else 1,
    )
    _populate_table(
        doc.tables[TABLE_MONTHLY],
        context.get("general_monthly", []),
        ("title", "start", "end", "status_text", "owner", "remark"),
        colors.get("monthly", []),
    )
    _populate_table(
        doc.tables[TABLE_QUARTERLY],
        context.get("general_quarterly", []),
        ("title", "start", "end", "status_text", "owner", "remark"),
        colors.get("quarterly", []),
    )
    _populate_table(
        doc.tables[TABLE_YEARLY],
        context.get("general_yearly", []),
        ("title", "start", "end", "status_text", "owner", "remark"),
        colors.get("yearly", []),
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
