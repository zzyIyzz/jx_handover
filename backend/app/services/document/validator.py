"""Fail-closed structural validation for generated handover DOCX files."""
from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document
from docx.oxml.ns import qn


MAIN_HEADINGS = (
    "一、基本信息",
    "二、设备变更情况",
    "三、重点工作完成情况",
    "四、需交接的工作",
    "五、对外委单位的考核",
    "六、定期工作完成情况",
)
SUB_HEADINGS = (
    "6.1月度定期工作",
    "6.2季度定期工作",
    "6.3年度定期工作",
)
BASIC_LABELS = (
    "交接开始时间",
    "交接截止时间",
    "交接班时间",
    "值班负责人",
    "临时值班负责人",
    "当班值班员",
)
TABLE_HEADERS = (
    None,
    ("序号", "工作内容", "开始时间", "结束时间", "完成人", "备注"),
    (
        "序号",
        "工作内容",
        "开始时间",
        "结束时间",
        "交接前责任人",
        "交接后责任人",
        "完成情况",
        "备注",
    ),
    ("序号", "外委单位", "工作内容", "考核情况", "备注"),
    ("序号", "工作内容", "开始时间", "结束时间", "完成情况", "完成人", "备注"),
    ("序号", "工作内容", "开始时间", "结束时间", "完成情况", "完成人", "备注"),
    ("序号", "工作内容", "开始时间", "结束时间", "完成情况", "完成人", "备注"),
)
EXPECTED_KEYS = ("important", "handover", "external", "monthly", "quarterly", "yearly")


class DocumentValidationError(RuntimeError):
    def __init__(self, errors: list[str]):
        self.errors = errors
        super().__init__("；".join(errors))


def validate_context(context: dict) -> dict:
    """Validate required business fields before touching a Word output file."""
    errors: list[str] = []
    if not str(context.get("station_name", "")).strip():
        errors.append("第一章缺少场站名称")
    for key, label in (
        ("start_date_cn", "交接开始时间"),
        ("end_date_cn", "交接截止时间"),
        ("handover_date_cn", "交接班时间"),
    ):
        if not str(context.get(key, "")).strip():
            errors.append(f"第一章缺少{label}")

    for index, row in enumerate(context.get("important_items", []), start=1):
        if not str(row.get("title", "")).strip():
            errors.append(f"第三章第 {index} 行缺少工作内容")
        if not str(row.get("completed_by", "")).strip():
            errors.append(f"第三章第 {index} 行缺少完成人")
    for index, row in enumerate(context.get("handover_items", []), start=1):
        if not str(row.get("title", "")).strip():
            errors.append(f"第四章第 {index} 行缺少工作内容")
        if not str(row.get("status_text", "")).strip():
            errors.append(f"第四章第 {index} 行缺少完成情况")
    for index, row in enumerate(context.get("external_rows", []), start=1):
        for key, label in (
            ("contractor", "外委单位"),
            ("work_content", "工作内容"),
            ("assessment", "考核情况"),
        ):
            if not str(row.get(key, "")).strip():
                errors.append(f"第五章第 {index} 行缺少{label}")
    if errors:
        raise DocumentValidationError(errors)
    return {"valid": True, "errors": [], "warnings": []}


def _normalize(value: str) -> str:
    return "".join(value.split())


def _expected_rows(key: str, count: int) -> int:
    if key == "external" and count == 0:
        return 3
    return max(1, count)


def validate_docx(path: Path, expected: dict[str, int] | None = None) -> dict:
    errors: list[str] = []
    warnings: list[str] = []
    try:
        with zipfile.ZipFile(path, "r") as archive:
            bad_member = archive.testzip()
            if bad_member:
                errors.append(f"DOCX ZIP 损坏：{bad_member}")
            if "word/document.xml" not in archive.namelist():
                errors.append("DOCX 缺少 word/document.xml")
    except (OSError, zipfile.BadZipFile) as exc:
        raise DocumentValidationError([f"无法读取 DOCX ZIP：{exc}"]) from exc

    if errors:
        raise DocumentValidationError(errors)
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        raise DocumentValidationError([f"Word 文档结构无法解析：{exc}"]) from exc

    body_texts = [paragraph.text.strip() for paragraph in doc.paragraphs]
    normalized_body = [_normalize(value) for value in body_texts]
    all_text = "\n".join(body_texts)

    if not body_texts or "交接班记录" not in body_texts[0]:
        errors.append("标题为空或不含“交接班记录”")
    if "\n" not in body_texts[0]:
        errors.append("标题未保持两行结构")
    for marker in ("{{", "}}", "场站名称交接班记录", "交接时段班次"):
        if marker in all_text:
            errors.append(f"存在未替换的模板标记：{marker}")

    positions: list[int] = []
    for heading in MAIN_HEADINGS:
        normalized = _normalize(heading)
        count = normalized_body.count(normalized)
        if count != 1:
            errors.append(f"主章节“{heading}”应出现 1 次，实际 {count} 次")
        else:
            positions.append(normalized_body.index(normalized))
    if len(positions) == len(MAIN_HEADINGS) and positions != sorted(positions):
        errors.append("一至六章顺序不正确")

    sub_positions: list[int] = []
    for heading in SUB_HEADINGS:
        normalized = _normalize(heading)
        count = normalized_body.count(normalized)
        if count != 1:
            errors.append(f"小节“{heading}”应出现 1 次，实际 {count} 次")
        else:
            sub_positions.append(normalized_body.index(normalized))
    if len(sub_positions) == len(SUB_HEADINGS) and sub_positions != sorted(sub_positions):
        errors.append("6.1、6.2、6.3 的顺序不正确")

    if len(doc.tables) != 7:
        errors.append(f"应有 7 张表（原六章表格及 6.3），实际 {len(doc.tables)} 张")
    else:
        labels = tuple(row.cells[0].text.strip() for row in doc.tables[0].rows)
        if labels != BASIC_LABELS:
            errors.append(f"第一章字段不一致：{labels}")
        for index, header in enumerate(TABLE_HEADERS[1:], start=1):
            actual = tuple(cell.text.strip() for cell in doc.tables[index].rows[0].cells)
            if actual != header:
                errors.append(f"第 {index + 1} 张表表头不一致：{actual}")
            header_properties = doc.tables[index].rows[0]._tr.get_or_add_trPr()
            if header_properties.find(qn("w:tblHeader")) is None:
                warnings.append(f"第 {index + 1} 张表未设置跨页重复表头")

        if expected is not None:
            for table_index, key in enumerate(EXPECTED_KEYS, start=1):
                expected_count = _expected_rows(key, int(expected.get(key, 0)))
                actual_count = len(doc.tables[table_index].rows) - 1
                if actual_count != expected_count:
                    errors.append(
                        f"{key} 数据行数应为 {expected_count}，实际 {actual_count}"
                    )
            if int(expected.get("external", 0)) == 0:
                numbers = [row.cells[0].text.strip() for row in doc.tables[3].rows[1:]]
                if numbers != ["1", "2", "3"]:
                    errors.append("第五章无数据时必须保留 1、2、3 三个占位行")

    if errors:
        raise DocumentValidationError(errors)
    return {"valid": True, "errors": [], "warnings": warnings}
