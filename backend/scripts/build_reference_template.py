"""Build the sanitized V0.3 Word template from the user's reference DOCX.

The reference file is deliberately supplied as a command-line argument and is
never copied into the repository.  Only the resulting, business-data-free
layout template is kept under ``backend/app/templates/word``.
"""
from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph


GENERIC_TITLE = "场站名称交接班记录\n（交接时段班次）"
GENERIC_DEVICE = "本班无设备变更。"


def _replace_paragraph_text(paragraph: Paragraph, value: str) -> None:
    """Replace text without discarding the reference run formatting."""
    if paragraph.runs:
        keep = paragraph.runs[0]._r
        for run in paragraph.runs[1:]:
            paragraph._p.remove(run._r)
    else:
        keep = paragraph.add_run()._r

    for child in list(keep):
        if child.tag != qn("w:rPr"):
            keep.remove(child)
    lines = value.split("\n")
    for index, line in enumerate(lines):
        text = OxmlElement("w:t")
        if line.startswith(" ") or line.endswith(" "):
            text.set(qn("xml:space"), "preserve")
        text.text = line
        keep.append(text)
        if index < len(lines) - 1:
            keep.append(OxmlElement("w:br"))


def _clear_cell(cell: _Cell) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = paragraphs[0]
        for extra in paragraphs[1:]:
            cell._tc.remove(extra._p)
    _replace_paragraph_text(paragraph, "")

    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:fill"), "FFFFFF")


def _clear_row(row: _Row) -> None:
    for cell in row.cells:
        _clear_cell(cell)


def _make_body_rows(table: Table, count: int, numbered: bool = False) -> None:
    if len(table.rows) < 2:
        raise ValueError("Reference table has no model body row")
    model_xml = deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    for index in range(count):
        table._tbl.append(deepcopy(model_xml))
        row = table.rows[-1]
        _clear_row(row)
        if numbered:
            _replace_paragraph_text(row.cells[0].paragraphs[0], str(index + 1))


def _append_yearly_section(doc: DocumentType) -> None:
    quarterly_heading = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("6.2")),
        None,
    )
    if quarterly_heading is None:
        raise ValueError("Reference document is missing the 6.2 heading")
    if len(doc.tables) != 6:
        raise ValueError(f"Expected six reference tables, got {len(doc.tables)}")

    body = doc._body._element
    section_properties = body.sectPr
    heading_xml = deepcopy(quarterly_heading._p)
    table_xml = deepcopy(doc.tables[5]._tbl)
    body.insert(body.index(section_properties), heading_xml)
    body.insert(body.index(section_properties), table_xml)

    yearly_heading = Paragraph(heading_xml, doc._body)
    yearly_table = Table(table_xml, doc._body)
    _replace_paragraph_text(yearly_heading, "6.3年度定期工作")
    _make_body_rows(yearly_table, 1)


def _remove_empty_body_paragraphs(doc: DocumentType) -> None:
    """Remove source-data spacer paragraphs that can strand a chapter title."""
    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip():
            paragraph._element.getparent().remove(paragraph._element)


def build_template(source: Path, output: Path) -> None:
    doc = Document(str(source))
    if len(doc.tables) != 6:
        raise ValueError(f"Expected six tables in reference DOCX, got {len(doc.tables)}")

    # Title and the only free-text block containing real equipment data.
    _replace_paragraph_text(doc.paragraphs[0], GENERIC_TITLE)
    section_two_index = next(
        i for i, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.startswith("二、")
    )
    _replace_paragraph_text(doc.paragraphs[section_two_index + 1], GENERIC_DEVICE)

    # Basic information: preserve row labels and layout, clear value cells.
    for row in doc.tables[0].rows:
        _clear_cell(row.cells[1])

    # Sections 3/4 keep a single style model row.  Section 5 intentionally
    # keeps the source document's three numbered blank placeholders.
    _make_body_rows(doc.tables[1], 1)
    _make_body_rows(doc.tables[2], 1)
    _make_body_rows(doc.tables[3], 3, numbered=True)
    _make_body_rows(doc.tables[4], 1)
    _make_body_rows(doc.tables[5], 1)
    _remove_empty_body_paragraphs(doc)
    _append_yearly_section(doc)

    properties = doc.core_properties
    properties.title = "江西片区智能交接班 V0.3 Word 母版"
    properties.subject = "脱敏版交接班记录版式母版"
    properties.author = "JXHandover"
    properties.last_modified_by = "JXHandover"
    properties.comments = "由用户原始 DOCX 脱敏生成；不含真实业务记录。"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    check = Document(str(output))
    if len(check.tables) != 7:
        raise RuntimeError("Sanitized template must contain seven tables")
    all_text = "\n".join(p.text for p in check.paragraphs)
    if GENERIC_TITLE.split("\n")[0] not in all_text or "6.3年度定期工作" not in all_text:
        raise RuntimeError("Sanitized template structure check failed")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_template(args.source.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
