# -*- coding: utf-8 -*-
"""验收检查脚本：读取命令行传入的 Word，打印结构与着色。"""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn

p = Path(sys.argv[1]).expanduser().resolve() if len(sys.argv) > 1 else None
if p is None or not p.is_file():
    raise SystemExit("用法：python check_docx.py <待检查的.docx>")
doc = Document(p)
for para in doc.paragraphs:
    if para.text.strip():
        print("P|", para.text[:70])
print("=== tables:", len(doc.tables))
for ti, t in enumerate(doc.tables):
    print(f"--- table {ti}: rows={len(t.rows)} cols={len(t.columns)}")
    for row in t.rows:
        tc_pr = row.cells[0]._tc.tcPr
        fill = ""
        if tc_pr is not None:
            shd = tc_pr.find(qn("w:shd"))
            if shd is not None:
                fill = shd.get(qn("w:fill")) or ""
        cells = " | ".join(c.text[:20] for c in row.cells[:4])
        print(f"  [{fill if fill else '白'}]", cells)
