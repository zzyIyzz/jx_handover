# -*- coding: utf-8 -*-
"""提取命令行传入 DOCX 的完整结构，供 Word 渲染格式对齐。"""
import os
from pathlib import Path
import sys
from docx import Document
from docx.oxml.ns import qn

OUT = os.path.join(os.environ.get("TEMP", "."), "docx2_structure.txt")


def cell_shading(cell):
    tcPr = cell._tc.tcPr
    if tcPr is not None:
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            return shd.get(qn("w:fill"))
    return None


def run_colors(run):
    """返回 (字体颜色, 是否加粗)"""
    color = None
    rPr = run._r.rPr
    if rPr is not None:
        c = rPr.find(qn("w:color"))
        if c is not None:
            color = c.get(qn("w:val"))
    return color, run.bold


def main():
    source = Path(sys.argv[1]).expanduser().resolve() if len(sys.argv) > 1 else None
    if source is None or not source.is_file():
        raise SystemExit("用法：python dump_docx2.py <参考模板.docx>")
    doc = Document(source)
    lines = []

    # 按 body 顺序遍历段落与表格
    body = doc.element.body
    tables = doc.tables
    tmap = {t._tbl: i for i, t in enumerate(tables)}
    ti = 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            text = p.text.strip()
            info = []
            for r in p.runs:
                c, b = run_colors(r)
                if c or b:
                    info.append(f"[{r.text!r} color={c} bold={b}]")
            style = p.style.name if p.style else ""
            lines.append(f"P|{text}|style={style}|" + " ".join(info))
        elif tag == "tbl":
            t = tables[ti]
            ti += 1
            lines.append(f"===== TABLE #{ti-1} rows={len(t.rows)} cols={len(t.columns)} =====")
            for ri, row in enumerate(t.rows):
                cells_desc = []
                for ci, cell in enumerate(row.cells):
                    txt = cell.text.replace("\n", "⏎").strip()
                    shd = cell_shading(cell)
                    # 收集单元格内字体颜色
                    fcolors = set()
                    for p in cell.paragraphs:
                        for r in p.runs:
                            c, b = run_colors(r)
                            if c:
                                fcolors.add(c)
                    desc = f"{txt}"
                    if shd and shd.lower() != "auto":
                        desc += f"«shd={shd}»"
                    if fcolors:
                        desc += f"«font={','.join(sorted(fcolors))}»"
                    cells_desc.append(desc)
                lines.append(f"  R{ri}: " + " | ".join(cells_desc))

    with open(OUT, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"OK -> {OUT} ({len(lines)} lines)")


if __name__ == "__main__":
    main()
