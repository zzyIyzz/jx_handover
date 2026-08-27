# -*- coding: utf-8 -*-
"""颜色专项验收：制造一个紧急 + 一个重点工单，渲染 V2，
核对 红底白字(FFA5A5+白字) / 黄底黑字(FFFE83) / 第三节归属与排序。
"""
import json
import os
import sys
import urllib.request

BASE = "http://127.0.0.1:8080/api"
OUT = os.path.join(os.environ.get("TEMP", "."), "color_check.txt")
L = []


def call(method, path, body=None):
    req = urllib.request.Request(
        BASE + path,
        data=json.dumps(body).encode("utf-8") if body is not None else None,
        method=method, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req) as r:
        return json.loads(r.read().decode("utf-8"))


def cell_shade(cell):
    from docx.oxml.ns import qn
    tc_pr = cell._tc.tcPr
    if tc_pr is not None:
        shd = tc_pr.find(qn("w:shd"))
        if shd is not None:
            return shd.get(qn("w:fill"))
    return None


def font_color(cell):
    runs = cell.paragraphs[0].runs
    return str(runs[0].font.color.rgb) if runs and runs[0].font.color and runs[0].font.color.rgb else "auto"


def main():
    batches = call("GET", "/handovers")
    bid = batches[0]["id"]
    det = call("GET", f"/handovers/{bid}")
    st = det["stations"][0]

    # 找一个未完成事项改紧急、一个已完成事项改重点
    target_u = next(i for i in st["items"] if i["status"] != "completed")
    target_c = next(i for i in st["items"] if i["status"] == "completed")
    call("PATCH", f"/handover-items/{target_u['id']}",
         {"revision": target_u["revision"], "priority": "urgent"})
    call("PATCH", f"/handover-items/{target_c['id']}",
         {"revision": target_c["revision"], "priority": "important"})
    L.append(f"紧急: {target_u['title']}")
    L.append(f"重点: {target_c['title']}")

    det = call("GET", f"/handovers/{bid}")
    st = det["stations"][0]
    for it in st["items"]:
        if it["priority"] in ("urgent", "important"):
            L.append(f"  {it['priority']}/{it['status']} -> {it['section']} "
                     f"color={it['color']} | {it['title'][:20]}")

    r = call("POST", f"/handovers/{bid}/render",
             {"station_meta_id": st["station_meta_id"]})
    L.append(f"渲染 V{r['version']} -> {r['docx_path']}")

    from docx import Document
    doc = Document(r["docx_path"])
    t1, t2 = doc.tables[1], doc.tables[2]
    L.append(f"第三节行数={len(t1.rows) - 1}（应含紧急+重点共 2 行）")
    for row in t1.rows[1:]:
        L.append(f"  三节行: {row.cells[1].text[:24]} | 底色="
                 f"{cell_shade(row.cells[0])} 字色={font_color(row.cells[1])}")
    cnt = {"FFA5A5": 0, "FFFE83": 0}
    for row in t2.rows[1:]:
        shd = cell_shade(row.cells[0])
        if shd in cnt:
            cnt[shd] += 1
            L.append(f"  四节行: {row.cells[1].text[:24]} | 底色={shd} "
                     f"字色={font_color(row.cells[1])}")
    L.append(f"第四节红底行={cnt['FFA5A5']} 黄底行={cnt['FFFE83']}")

    with open(OUT, "w", encoding="utf-8") as f:
        f.write("\n".join(L))
    print("OK ->", OUT)


if __name__ == "__main__":
    main()
