"""生成 Word 模板 handover_v1.docx（docxtpl Jinja2 标签）。

结构完全对齐现有修水交接班记录的六节格式。
所有标签写为单个 run，避免 docxtpl 解析拆行问题。
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt  # noqa: E402

from app.config import WORD_TEMPLATE  # noqa: E402

FONT = "宋体"
# 参考 docx 实测色值：表头灰 / 基本信息标签列浅灰
SHD_HEADER = "D9D9D9"
SHD_LABEL = "F2F2F2"


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_font(run, size=10.5, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_para(doc, text, size=10.5, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    _set_font(p.add_run(text), size, bold)
    return p


def add_table(doc, headers: list[str], tag_rows: list[list[str]],
              widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _set_font(cell.paragraphs[0].add_run(h), 10.5, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade(cell, SHD_HEADER)
    for row_texts in tag_rows:
        row = table.add_row()
        for i, t in enumerate(row_texts):
            cell = row.cells[i]
            cell.text = ""
            _set_font(cell.paragraphs[0].add_run(t))
    return table


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # 标题（与参考 docx 一致：单段落 “场站名交接班记录（区间班次）”）
    add_para(doc, "{{ station_name }}交接班记录（{{ period_cn }}班次）",
             16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 一、基本信息（表 0）
    add_para(doc, "一、基本信息", 12, bold=True)
    basic = [
        ("交接开始时间", "{{ start_date_cn }}"),
        ("交接截止时间", "{{ end_date_cn }}"),
        ("交接班时间", "{{ handover_date_cn }}"),
        ("值班负责人", "{{ duty_leader }}"),
        ("临时值班负责人", "{{ temp_leader }}"),
        ("当班值班员", "{{ operators }}"),
    ]
    t0 = doc.add_table(rows=len(basic), cols=2)
    t0.style = "Table Grid"
    for i, (k, v) in enumerate(basic):
        c0, c1 = t0.rows[i].cells
        c0.text = ""
        _set_font(c0.paragraphs[0].add_run(k), bold=True)
        _shade(c0, SHD_LABEL)
        c1.text = ""
        _set_font(c1.paragraphs[0].add_run(v))

    # 二、设备变更情况（{%p %} 标签必须各自独占一个段落）
    add_para(doc, "二、设备变更情况", 12, bold=True)
    add_para(doc, "{%p for c in device_changes %}")
    add_para(doc, "{{ loop.index }}、{{ c }}")
    add_para(doc, "{%p endfor %}")
    add_para(doc, "{%p if not device_changes %}")
    add_para(doc, "本班无设备变更。")
    add_para(doc, "{%p endif %}")

    # 三、重点工作完成情况（表 1）
    add_para(doc, "三、重点工作完成情况", 12, bold=True)
    add_para(doc, "（紧急＝红、重点＝黄的重点/临时工作归入本节；"
                  "无则写“本班无紧急/重点工作”。）", 9)
    add_table(
        doc,
        ["序号", "工作内容", "开始时间", "结束时间", "完成人", "备注"],
        [
            ["{%tr for item in important_items %}"],
            ["{{ loop.index }}", "{{ item.title }}", "{{ item.start }}",
             "{{ item.end }}", "{{ item.owner }}", "{{ item.remark }}"],
            ["{%tr endfor %}"],
            ["{%tr if not important_items %}"],
            ["1", "本班无紧急/重点工作", "—", "—", "—", ""],
            ["{%tr endif %}"],
        ],
    )

    # 四、需交接的工作（表 2）
    add_para(doc, "四、需交接的工作", 12, bold=True)
    add_para(doc, "（第三、六节之外的未完成/进行中工作归入本节，移交下一班；"
                  "第三、六节中的工作不重复列入。紧急工作标红，重点工作标黄）", 9)
    add_table(
        doc,
        ["序号", "工作内容", "开始时间", "结束时间", "交接前责任人",
         "交接后责任人", "完成情况", "备注"],
        [
            ["{%tr for item in handover_items %}"],
            ["{{ loop.index }}", "{{ item.title }}", "{{ item.start }}",
             "{{ item.end }}", "{{ item.prev_owner }}",
             "{{ item.next_owner }}", "{{ item.status_text }}",
             "{{ item.remark }}"],
            ["{%tr endfor %}"],
            ["{%tr if not handover_items %}"],
            ["1", "本班无需交接工作", "—", "—", "—", "—", "—", ""],
            ["{%tr endif %}"],
        ],
    )

    # 五、对外委单位的考核（表 3）
    add_para(doc, "五、对外委单位的考核", 12, bold=True)
    add_table(
        doc,
        ["序号", "外委单位", "工作内容", "考核情况", "备注"],
        [
            ["{%tr for row in external_rows %}"],
            ["{{ row.no }}", "", "", "", ""],
            ["{%tr endfor %}"],
        ],
    )

    # 六、定期工作完成情况（内置模板库自动拉取，程序判定颜色与排序）
    add_para(doc, "六、定期工作完成情况", 12, bold=True)
    add_para(doc, "（本交接时段到期的月度/季度/年度工作由程序自动全部列入本节："
                  "已完成＝绿，在期限内未完成＝白，超期限未完成＝红。"
                  "排序规则：先未完成，后已完成。）", 9)
    general_tables = [
        ("6.1月度定期工作", "general_monthly", "本班无到期月度定期工作"),
        ("6.2季度定期工作", "general_quarterly", "本班无到期季度定期工作"),
        ("6.3年度定期工作", "general_yearly", "本班无到期年度定期工作"),
    ]
    for heading, var, empty_text in general_tables:
        add_para(doc, heading, 11, bold=True)
        add_table(
            doc,
            ["序号", "工作内容", "开始时间", "结束时间", "完成情况",
             "完成人", "备注"],
            [
                ["{%tr for item in " + var + " %}"],
                ["{{ loop.index }}", "{{ item.title }}", "{{ item.start }}",
                 "{{ item.end }}", "{{ item.status_text }}",
                 "{{ item.owner }}", "{{ item.remark }}"],
                ["{%tr endfor %}"],
                ["{%tr if not " + var + " %}"],
                ["1", empty_text, "—", "—", "—", "—", ""],
                ["{%tr endif %}"],
            ],
        )

    WORD_TEMPLATE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(WORD_TEMPLATE))
    print(f"模板已生成: {WORD_TEMPLATE}")


if __name__ == "__main__":
    build()
