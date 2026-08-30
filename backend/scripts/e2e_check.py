# -*- coding: utf-8 -*-
"""端到端验收：修水眉毛山 2026.8.14~8.23 演示数据 -> 正式 Word。
逐项核对：模板项无多余无遗漏、三/四/六节排序、第五节渲染、颜色。
结果写入 %TEMP%\\e2e_report.txt（UTF-8）。
"""
import json
import os
import sys
import urllib.request

BASE = "http://127.0.0.1:8765/api"
OUT = os.path.join(os.environ.get("TEMP", "."), "e2e_report.txt")
L = []


def call(method, path, body=None):
    req = urllib.request.Request(
        BASE + path,
        data=json.dumps(body).encode("utf-8") if body is not None else None,
        method=method,
        headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req) as r:
        return json.loads(r.read().decode("utf-8"))


def cell_shade(cell):
    from docx.oxml.ns import qn
    tc_pr = cell._tc.tcPr
    if tc_pr is not None:
        shd = tc_pr.find(qn("w:shd"))
        if shd is not None:
            return shd.get(qn("w:fill"))
    return None


def main():
    stations = call("GET", "/stations")
    xs = next(s for s in stations if s["code"] == "XS_MMS")
    L.append(f"场站: {xs['code']} {xs['name']} id={xs['id']}")

    # 1. 新建班次
    batch = call("POST", "/handovers", {
        "start_date": "2026-08-14", "end_date": "2026-08-23",
        "handover_date": "2026-08-23", "station_ids": [xs["id"]],
        "meta_overrides": {str(xs["id"]): {
            "duty_leader": "熊思奇", "temp_leader": "无",
            "operators": ["盛林"]}}})
    bid = batch["id"]
    L.append(f"班次已创建: {bid}")

    # 2. 详情核对
    det = call("GET", f"/handovers/{bid}")
    st = det["stations"][0]
    L.append(f"专业事项 {len(st['items'])} 条:")
    for it in st["items"]:
        L.append(f"  [{it['priority']}/{it['status']}] -> {it['section']} | "
                 f"{it['color']} | {it['title']}")
    gen = st["general"]
    L.append("定期工作: 月度 %d / 季度 %d / 年度 %d"
             % (len(gen["monthly"]), len(gen["quarterly"]),
                len(gen["yearly"])))
    expect = {"monthly": 45, "quarterly": 2, "yearly": 5}
    for cat, n in expect.items():
        got = len(gen[cat])
        L.append(f"  验收: {cat} 应为 {n} 项 -> 实际 {got} "
                 f"{'OK' if got == n else 'FAIL'}")

    # 排序核对：各节先未完成后已完成
    def order_ok(rows):
        seen_done = False
        for r in rows:
            done = r["status"] == "completed"
            if seen_done and not done:
                return False
            seen_done = seen_done or done
        return True

    for cat, rows in gen.items():
        L.append(f"  排序({cat} 未完成在前): "
                 + ("OK" if order_ok(rows) else "FAIL"))
        reds = [r["title"] for r in rows if r["color"] == "red"]
        L.append(f"    超期红: {reds if reds else '无'}")

    # 3. 全部复核确认
    pend = [i for i in st["items"] if i["review_status"] == "pending"]
    for it in pend:
        call("POST", f"/handover-items/{it['id']}/approve",
             {"revision": it["revision"]})
    L.append(f"已确认 {len(pend)} 条待复核事项")

    # 4. 设备变更
    for c in ("#1SVG为检修状态", "F08风机、F13风机箱变接地保护已退出"):
        call("POST", f"/handovers/{bid}/device-changes",
             {"station_meta_id": st["station_meta_id"], "content": c})

    # 5. 渲染正式 Word
    r = call("POST", f"/handovers/{bid}/render",
             {"station_meta_id": st["station_meta_id"]})
    L.append(f"Word 已生成: V{r['version']} -> {r['docx_path']}")

    # 6. 解析 docx 校验结构与颜色
    from docx import Document
    doc = Document(r["docx_path"])
    L.append(f"\nDOCX 段落/表格核对（表格数={len(doc.tables)}，期望 7）:")
    heads = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and (t.startswith(("一、", "二、", "三、", "四、", "五、",
                                "六、", "6.")) or "（" in t and len(t) < 120):
            heads.append(t[:80])
    for h in heads:
        L.append("  " + h)

    names = ["基本信息", "重点工作", "需交接", "外委考核",
             "月度定期", "季度定期", "年度定期"]
    # 完成情况所在列：表 1 看备注列(5)，表 2 完成情况列(6)，定期表列(4)
    status_idx = {1: 5, 2: 6, 4: 4, 5: 4, 6: 4}
    for ti, table in enumerate(doc.tables):
        rows = table.rows
        first_shd = cell_shade(rows[0].cells[0])
        cnt = {"C6EFCE": 0, "FFA5A5": 0, "FFFE83": 0, "D9D9D9": 0}
        seq = []
        for row in rows[1:]:
            shd = cell_shade(row.cells[0])
            if shd in cnt:
                cnt[shd] += 1
            if ti in status_idx:
                text = row.cells[status_idx[ti]].text.strip()
                done = ("已完成" in text) and ("未完成" not in text)
                seq.append("完" if done else "未")
        L.append(f"  表{ti}[{names[ti]}] rows={len(rows)} "
                 f"表头底色={first_shd} 数据行: 绿{cnt['C6EFCE']} "
                 f"红{cnt['FFA5A5']} 黄{cnt['FFFE83']}")
        if seq:
            ok = not any(a == "完" and b == "未"
                         for a, b in zip(seq, seq[1:]))
            L.append(f"    排序序列({''.join(seq)[:60]}…) "
                     + ("OK" if ok else "FAIL"))

    # 外委考核表头核对
    t3 = doc.tables[3]
    hdr = [c.text for c in t3.rows[0].cells]
    L.append(f"  第五节表头: {hdr} 占位行数={len(t3.rows) - 1}")

    # 紧急行字体颜色核对（红底白字）
    t1 = doc.tables[1]
    for row in t1.rows[1:]:
        if cell_shade(row.cells[0]) == "FFA5A5":
            runs = row.cells[1].paragraphs[0].runs
            fc = runs[0].font.color.rgb if runs else None
            L.append(f"  红底行字体色={fc}（期望 FFFFFF）: "
                     + ("OK" if str(fc) == "FFFFFF" else "FAIL"))
            break

    with open(OUT, "w", encoding="utf-8") as f:
        f.write("\n".join(L))
    print("OK ->", OUT)


if __name__ == "__main__":
    main()
